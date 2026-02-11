from docx import Document
import os
import uuid


class WordReportService:
    def generate_docx(self, text: str, title: str = "AI Analysis Report") -> str:
        document = Document()

        # Заголовок
        document.add_heading(title, level=1)

        # Статистика
        document.add_paragraph(f"Characters: {len(text)}")
        document.add_paragraph(f"Words: {len(text.split())}")

        document.add_paragraph("")  # порожній рядок

        # Основний текст
        document.add_paragraph(text)

        # Тимчасова папка
        temp_dir = "temp_reports"
        os.makedirs(temp_dir, exist_ok=True)

        filename = os.path.join(
            temp_dir,
            f"report_{uuid.uuid4()}.docx"
        )

        document.save(filename)
        return filename


word_report_service = WordReportService()
